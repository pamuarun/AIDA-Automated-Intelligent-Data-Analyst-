import io
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Inches


class ReportingAgent:
    """Generates PDF and DOCX reports using PNG plot bytes."""

    def generate_reports(self, original_df, cleaned_df, eda_output, cleaning_summary, insights):

        # Extract plot bytes from EDA output if available
        plot_bytes_map = eda_output.get("plot_bytes", {}) or {}

        # -------------------------------------------------------------------
        # PDF REPORT
        # -------------------------------------------------------------------
        pdf_buf = io.BytesIO()
        c = canvas.Canvas(pdf_buf, pagesize=letter)
        width, height = letter

        c.setFont("Helvetica-Bold", 14)
        c.drawString(40, height - 40, "Automated Data Analyst Report")

        c.setFont("Helvetica", 10)

        y = height - 80

        # Executive Summary
        exec_sum = insights.get("executive_summary") if isinstance(insights, dict) else None

        if exec_sum:
            c.drawString(40, y, "Executive Summary:")
            y -= 20

            text = c.beginText(40, y)
            for line in exec_sum.split("\n"):
                text.textLine(line)
            c.drawText(text)
            y -= 120

        # Dataset Shapes
        c.drawString(40, y, f"Original Shape: {original_df.shape}")
        c.drawString(300, y, f"Cleaned Shape: {cleaned_df.shape}")
        y -= 30

        # Cleaning Summary
        c.drawString(40, y, "Cleaning Summary:")
        y -= 20

        if cleaning_summary:
            if cleaning_summary.get("duplicates_removed") is not None:
                c.drawString(60, y, f"Duplicates Removed: {cleaning_summary['duplicates_removed']}")
                y -= 16

            if cleaning_summary.get("imputation"):
                cols = ", ".join(cleaning_summary["imputation"].keys())
                c.drawString(60, y, f"Imputed Columns: {cols}")
                y -= 16

        # -------------------------------------------------------------------
        # Insert Correlation Heatmap (PNG bytes)
        # -------------------------------------------------------------------
        corr_png_bytes = plot_bytes_map.get("correlation_heatmap")

        if corr_png_bytes:
            try:
                img_buf = io.BytesIO(corr_png_bytes)

                c.drawImage(img_buf, 40, y - 220, width=500, height=200)
                y -= 260

            except Exception as e:
                print("PDF heatmap error:", e)

        c.showPage()
        c.save()
        pdf_bytes = pdf_buf.getvalue()

        # -------------------------------------------------------------------
        # DOCX REPORT
        # -------------------------------------------------------------------
        doc = Document()
        doc.add_heading("Automated Data Analyst Report", level=1)

        if exec_sum:
            doc.add_heading("Executive Summary", level=2)
            doc.add_paragraph(exec_sum)

        doc.add_heading("Cleaning Summary", level=2)
        if cleaning_summary:
            if cleaning_summary.get("duplicates_removed") is not None:
                doc.add_paragraph(f"Duplicates removed: {cleaning_summary['duplicates_removed']}")

            if cleaning_summary.get("imputation"):
                doc.add_paragraph(
                    f"Imputed columns: {', '.join(cleaning_summary['imputation'].keys())}"
                )

        # Heatmap in DOCX
        if corr_png_bytes:
            try:
                img_buf = io.BytesIO(corr_png_bytes)
                doc.add_heading("Correlation Heatmap", level=2)
                doc.add_picture(img_buf, width=Inches(6))

            except Exception as e:
                print("DOCX heatmap error:", e)

        # Export DOCX
        docx_buf = io.BytesIO()
        doc.save(docx_buf)

        return {
            "pdf_bytes": pdf_bytes,
            "docx_bytes": docx_buf.getvalue()
        }
